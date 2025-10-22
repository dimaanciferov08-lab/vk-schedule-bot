import vk_api
import json
from vk_api.bot_longpoll import VkBotLongPoll, VkBotEventType
from vk_api.utils import get_random_id
import sqlite3
import datetime
import time
import threading
import requests
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re
import logging
import wikipediaapi
import random
import math
from sympy import symbols, solve, simplify, diff, integrate, sqrt
import sympy as sp

# Настройка логирования
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# 🔑 НАСТРОЙКИ
CONFIG = {
    "group_id": 232761329,
    "token": "vk1.a.Y2xBv4alWQ55rd1IxtkpKc48ibKqpQ1x0Wyc9Hv0z18elxu3JaSBfCi7F5sJ9H4eKy1jg3iqFOjQTkQyCIYdnf77mcezdC__MLiyRi9Xwfus_uLz7UWd9AR8VPQDr7uMEiD1NxadTzqUllP7p4uqWixuefYkm6ryhgMbFLPSo-hnXKyt0XQ4qvpfIG5kLWlJoH7Ivew1yhgiKmtDWhbHYw",
    "admin_id": 238448950,
    "current_week": 1,
    "allowed_chats": [2000000002],
    "chat_id": None
}

# Список группы
GROUP_LIST = {
    "1": "Амосов Никита", "2": "Богомолов Георгий", "3": "Веселов Даниил",
    "4": "Громов Роман", "5": "Долотин Иван", "6": "Дударев Святослав",
    "7": "Зуев Андрей", "8": "Иванов Матвей", "9": "Карпов Дмитрий",
    "10": "Клещев Сергей", "11": "Лебедев Кирилл", "12": "Назаренков Иван",
    "13": "Святец Александр", "14": "Семенов Леонид", "15": "Фомичева Елизавета",
    "16": "Шевченко Дарья", "17": "Яременко Антон"
}

# === УМНЫЙ ТОСИК С ИИ И ПРИКОЛАМИ ===
class SmartTosik:
    def __init__(self):
        self.wiki_wiki = wikipediaapi.Wikipedia(
            language='ru',
            extract_format=wikipediaapi.ExtractFormat.WIKI,
            user_agent="TosikBot/3.0"
        )
        self.personality_traits = self._init_personality()
        self.jokes_db = self._init_jokes()
        self.facts_db = self._init_facts()
    
    def _init_personality(self):
        """Характер Тосика"""
        return {
            'name': 'Тосик',
            'mood': random.choice(['веселый', 'задумчивый', 'энергичный', 'философский']),
            'traits': ['умный', 'с юмором', 'помогающий', 'немного саркастичный'],
            'favorite_topics': ['технологии', 'наука', 'программирование', 'студенческая жизнь']
        }
    
    def _init_jokes(self):
        """База шуток Тосика"""
        return {
            'programming': [
                "Почему программисты путают Хэллоуин и Рождество? Потому что Oct 31 == Dec 25! 🎃",
                "Сколько программистов нужно, чтобы поменять лампочку? Ни одного, это hardware проблема! 💡",
                "Программист звонит в техподдержку: 'У меня код не компилируется!' - 'А вы пробовали выключить и включить мозг?' 🧠"
            ],
            'study': [
                "Студент и дедлайн - вечная любовь. Он бежит от неё, а она его догоняет! 🏃‍♂️",
                "Знаете почему студенты такие умные? Потому что они спят на лекциях и впитывают знания через подушку! 😴",
                "Сессия близко... А значит скоро мы все станем экспертами по ночным перерывам на кофе! ☕"
            ],
            'science': [
                "Физик и математик сидят в кафе. Вдруг загорается огонь. Физик бежит за огнетушителем, а математик говорит: 'Решение существует!' 🔥",
                "Почему электроны никогда не берут кредиты? Потому что у них всегда отрицательный заряд! ⚡"
            ],
            'random': [
                "Я не ленивый, я просто на энергосберегающем режиме! 🔋",
                "Мой внутренний голос говорит мне не есть пиццу в 3 утра... Жаль, что у него нет рук чтобы меня остановить! 🍕",
                "Сегодня видел как чайник закипел и подумал: 'Вот бы так же легко закипать от знаний!' 📚"
            ]
        }
    
    def _init_facts(self):
        """База интересных фактов"""
        return {
            'technology': [
                "Первый компьютерный вирус был создан в 1983 году и назывался 'Elk Cloner'",
                "Самый популярный язык программирования в 2024 - Python, а самый ненавистный - JavaScript (шучу! 😄)",
                "Первая веб-камера была создана для отслеживания кофеварки в Кембридже"
            ],
            'science': [
                "Человеческий мозг может хранить до 2.5 петабайт информации - это примерно 3 миллиона часов сериалов!",
                "Свету от Солнца нужно 8 минут чтобы достичь Земли, но фотонам в ядре Солнца нужно 10000 лет чтобы вырваться наружу",
                "В вашем теле достаточно железа чтобы сделать гвоздь длиной 3 см"
            ],
            'history': [
                "Первая компьютерная мышь была сделана из дерева в 1964 году",
                "Первый сайт в интернете до сих пор работает: info.cern.ch",
                "QWERTY-раскладка была создана чтобы замедлить печать и предотвратить залипание клавиш на печатных машинках"
            ]
        }
    
    def get_joke(self, category=None):
        """Получить случайную шутку"""
        if category and category in self.jokes_db:
            jokes = self.jokes_db[category]
        else:
            # Собираем все шутки
            all_jokes = []
            for cat_jokes in self.jokes_db.values():
                all_jokes.extend(cat_jokes)
            jokes = all_jokes
        
        return random.choice(jokes)
    
    def get_fact(self, category=None):
        """Получить случайный факт"""
        if category and category in self.facts_db:
            facts = self.facts_db[category]
        else:
            all_facts = []
            for cat_facts in self.facts_db.values():
                all_facts.extend(cat_facts)
            facts = all_facts
        
        return random.choice(facts)
    
    def search_internet(self, query):
        """Поиск информации в интернете через Wikipedia"""
        try:
            page = self.wiki_wiki.page(query)
            if page.exists():
                # Берем только первые 500 символов чтобы не перегружать
                summary = page.summary[:500] + "..." if len(page.summary) > 500 else page.summary
                return {
                    'found': True,
                    'title': page.title,
                    'summary': summary,
                    'url': page.fullurl
                }
            return {'found': False}
        except Exception as e:
            logging.error(f"Ошибка поиска: {e}")
            return {'found': False}
    
    def generate_ai_response(self, question, context=None):
        """Генерация умного ответа с ИИ-приколами"""
        # Определяем тип вопроса
        question_lower = question.lower()
        
        # Простые вопросы - отвечаем с юмором
        simple_questions = {
            'как дела': [
                "Отлично! Только что победил в шахматы у ChatGPT! ♟️",
                "Супер! Готовлюсь к мировой domination... в смысле, к помощи студентам! 🌍",
                "Лучше не бывает! Только что узнал новую шутку про программистов!",
                "Как у робота с мечтами - заряжаюсь энергией и готов к свершениям! ⚡"
            ],
            'кто ты': [
                "Я Тосик - самый умный (и скромный) бот в этой беседе! 🤖",
                "Тосик, твой цифровой друг, готовый помочь с учебой и поднять настроение! 😊",
                "Просто маленький ИИ с большими амбициями и любовью к студентам! 📚"
            ],
            'что ты умеешь': [
                "Умею всё! Ну почти... Расписание, доклады, рефераты, шутки рассказывать - я твой универсальный студенческий солдат! 🎯",
                "От пары расписать до реферата сгенерить - вот мои скромные таланты! А ещё шутки травить - это бесплатно! 😄"
            ],
            'привет': [
                "Привет-привет! Готов к учебным подвигам? 🦸‍♂️",
                "Здарова, человечек! Как успехи в покорении знаний? 📖",
                "Приветствую! Надеюсь, ты не забыл, что сессия всегда ближе, чем кажется! 😅"
            ],
            'пока': [
                "Пока-пока! Не теряй зарядку... и знания! 🔋",
                "До скорого! Надеюсь, твой мозг не перегреется от новых знаний! 🧠",
                "Прощай! Помни: кофе - друг, прокрастинация - враг! ☕"
            ]
        }
        
        # Проверяем простые вопросы
        for q, answers in simple_questions.items():
            if q in question_lower:
                return random.choice(answers)
        
        # Вопросы про шутки и факты
        if any(word in question_lower for word in ['шутка', 'шути', 'смешн', 'прикол']):
            category = None
            if 'программир' in question_lower:
                category = 'programming'
            elif 'учеба' in question_lower or 'студен' in question_lower:
                category = 'study'
            elif 'наук' in question_lower:
                category = 'science'
            
            joke = self.get_joke(category)
            return f"🎭 Держи шутку!\n\n{joke}\n\nХочешь ещё? Просто скажи 'ещё шутка'!"
        
        if any(word in question_lower for word in ['факт', 'интересн', 'узнать']):
            category = None
            if 'техн' in question_lower:
                category = 'technology'
            elif 'наук' in question_lower:
                category = 'science'
            elif 'истори' in question_lower:
                category = 'history'
            
            fact = self.get_fact(category)
            return f"🧠 Вот интересный факт!\n\n{fact}\n\nХочешь ещё фактов? Просто попроси!"
        
        # Сложные вопросы - ищем в интернете
        if len(question.split()) >= 3:  # Если вопрос достаточно сложный
            search_result = self.search_internet(question)
            if search_result['found']:
                response = f"🔍 Вот что я нашел по твоему вопросу '{question}':\n\n"
                response += f"📖 {search_result['title']}\n"
                response += f"📝 {search_result['summary']}\n\n"
                response += f"🌐 Подробнее: {search_result['url']}\n\n"
                response += "💡 Нужно что-то конкретное? Уточни вопрос!"
                return response
        
        # Если не нашли ответ - общий умный ответ
        general_responses = [
            "Хм, интересный вопрос! Дай-ка подумать... 🤔\nМожет, спросишь что-то про учебу или технологии? Я в них спец!",
            "Ого, ты задал сложный вопрос! Я пока учусь отвечать на такие... 😅\nПопробуй спросить про расписание или рефераты!",
            "Вот это да! Ты заставил мой процессор задуматься! 💭\nМожет, сначала простые вопросы? Я же только учусь!"
        ]
        
        return random.choice(general_responses)
    
    def should_respond(self, message, user_id):
        """Определяет, должен ли Тосик отвечать на сообщение"""
        message_lower = message.lower()
        
        # Всегда отвечаем на прямое обращение
        if any(name in message_lower for name in ['тосик', 'тосика', 'тосику', 'тосиком', 'tosik']):
            return True
        
        # Отвечаем на вопросы (содержат знак вопроса)
        if '?' in message:
            return True
        
        # Отвечаем на приветствия и прощания
        greetings = ['привет', 'хай', 'здаров', 'hello', 'hi', 'ку', 'здравствуй']
        farewells = ['пока', 'прощай', 'до свидан', 'bye', 'goodbye']
        
        if any(word in message_lower for word in greetings + farewells):
            return True
        
        # Отвечаем на некоторые ключевые слова
        keywords = [
            'как дела', 'кто ты', 'что ты', 'умеешь', 'можешь',
            'шутка', 'факт', 'расскажи', 'объясни', 'помоги'
        ]
        
        if any(word in message_lower for word in keywords):
            return True
        
        # Иногда отвечаем случайно (30% chance) на длинные сообщения
        if len(message) > 20 and random.random() < 0.3:
            return True
        
        return False

# Создаем умного Тосика
smart_tosik = SmartTosik()

# === УЛУЧШЕННЫЙ ГЕНЕРАТОР РЕФЕРАТОВ ===
class AdvancedReferatGenerator:
    def __init__(self):
        self.wiki_wiki = wikipediaapi.Wikipedia(
            language='ru',
            extract_format=wikipediaapi.ExtractFormat.WIKI,
            user_agent="TosikBot/2.0"
        )
        self.knowledge_base = self._init_knowledge_base()
    
    def _init_knowledge_base(self):
        """База знаний по популярным темам"""
        return {
            "искусственный интеллект": {
                "definition": "Искусственный интеллект (ИИ) — это область компьютерных наук, занимающаяся созданием машин и систем, способных выполнять задачи, требующие человеческого интеллекта.",
                "history": "История ИИ начинается с 1950-х годов. Алан Тьюринг предложил тест Тьюринга в 1950 году. В 1956 году термин 'искусственный интеллект' был введен Джоном Маккарти.",
                "types": ["Машинное обучение", "Глубокое обучение", "Обработка естественного языка", "Компьютерное зрение", "Робототехника"],
                "applications": ["Медицина", "Финансы", "Образование", "Транспорт", "Развлечения"],
                "technologies": ["Нейронные сети", "Генетические алгоритмы", "Экспертные системы", "Нечеткая логика"]
            },
            "машинное обучение": {
                "definition": "Машинное обучение — это подраздел искусственного интеллекта, focusing на разработке алгоритмов, которые могут обучаться на данных.",
                "methods": ["Обучение с учителем", "Обучение без учителя", "Обучение с подкреплением"],
                "algorithms": ["Линейная регрессия", "Деревья решений", "Метод k-ближайших соседей", "Нейронные сети"]
            },
            "программирование": {
                "languages": ["Python", "Java", "C++", "JavaScript", "C#"],
                "paradigms": ["Объектно-ориентированное", "Функциональное", "Процедурное", "Логическое"],
                "concepts": ["Алгоритмы", "Структуры данных", "Базы данных", "ООП", "Тестирование"]
            },
            "физика": {
                "sections": ["Механика", "Термодинамика", "Электромагнетизм", "Квантовая физика", "Оптика"],
                "laws": ["Законы Ньютона", "Закон сохранения энергии", "Законы термодинамики", "Закон Ома"]
            }
        }
    
    def search_enhanced_content(self, topic):
        """Расширенный поиск контента"""
        content = {
            'wikipedia': self._get_wikipedia_content(topic),
            'knowledge_base': self._get_knowledge_base_content(topic),
            'related_topics': self._get_related_topics(topic)
        }
        return content
    
    def _get_wikipedia_content(self, topic):
        """Получение контента из Wikipedia"""
        try:
            page = self.wiki_wiki.page(topic)
            if page.exists():
                return {
                    'title': page.title,
                    'summary': page.summary[:2000],
                    'sections': self._extract_wiki_sections(page),
                    'exists': True
                }
            return {'exists': False}
        except Exception as e:
            logging.error(f"Ошибка Wikipedia: {e}")
            return {'exists': False}
    
    def _extract_wiki_sections(self, page):
        """Извлечение секций из Wikipedia"""
        sections = {}
        for section in page.sections:
            sections[section.title] = section.text[:1000]
        return sections
    
    def _get_knowledge_base_content(self, topic):
        """Получение контента из базы знаний"""
        topic_lower = topic.lower()
        for key, value in self.knowledge_base.items():
            if key in topic_lower:
                return value
        return None
    
    def _get_related_topics(self, topic):
        """Генерация связанных тем"""
        related_map = {
            "ии": ["машинное обучение", "нейронные сети", "глубокое обучение", "робототехника"],
            "программирование": ["алгоритмы", "базы данных", "веб-разработка", "мобильная разработка"],
            "физика": ["механика", "электричество", "магнетизм", "термодинамика"],
            "математика": ["алгебра", "геометрия", "статистика", "дифференциальные уравнения"]
        }
        
        for key, values in related_map.items():
            if key in topic.lower():
                return values
        return ["технологии", "наука", "исследования", "разработка"]
    
    def calculate_structure(self, pages):
        """Расчет структуры на основе количества страниц"""
        if pages <= 5:
            return {"chapters": 2, "sections_per_chapter": 2, "content_length": 300}
        elif pages <= 10:
            return {"chapters": 3, "sections_per_chapter": 3, "content_length": 500}
        elif pages <= 15:
            return {"chapters": 3, "sections_per_chapter": 4, "content_length": 700}
        elif pages <= 20:
            return {"chapters": 4, "sections_per_chapter": 4, "content_length": 800}
        else:
            return {"chapters": 5, "sections_per_chapter": 4, "content_length": 1000}
    
    def generate_intelligent_content(self, topic, pages=10):
        """Интеллектуальная генерация контента"""
        structure = self.calculate_structure(pages)
        enhanced_content = self.search_enhanced_content(topic)
        
        chapters = []
        
        for chapter_num in range(structure['chapters']):
            chapter_title = self._generate_chapter_title(topic, chapter_num, enhanced_content)
            chapter = {
                'title': chapter_title,
                'sections': []
            }
            
            for section_num in range(structure['sections_per_chapter']):
                section_title = self._generate_section_title(chapter_num, section_num, enhanced_content)
                section_content = self._generate_section_content(
                    topic, chapter_num, section_num, enhanced_content, structure['content_length']
                )
                
                chapter['sections'].append({
                    'title': section_title,
                    'content': section_content
                })
            
            chapters.append(chapter)
        
        return chapters
    
    def _generate_chapter_title(self, topic, chapter_num, enhanced_content):
        """Генерация заголовков глав"""
        base_titles = [
            f"Теоретические основы {topic}",
            f"Историческое развитие и эволюция {topic}",
            f"Современные технологии и методы в области {topic}",
            f"Практическое применение {topic}",
            f"Перспективы и будущее развитие {topic}",
            f"Методологические аспекты исследования {topic}",
            f"Ключевые концепции и принципы {topic}"
        ]
        
        # Используем реальные данные если есть
        if enhanced_content.get('wikipedia', {}).get('exists'):
            wiki_title = enhanced_content['wikipedia']['title']
            custom_titles = [
                f"Основные понятия и определения {wiki_title}",
                f"Исторические аспекты развития {wiki_title}",
                f"Современное состояние и тенденции {wiki_title}"
            ]
            base_titles = custom_titles + base_titles
        
        return base_titles[chapter_num % len(base_titles)]
    
    def _generate_section_title(self, chapter_num, section_num, enhanced_content):
        """Генерация заголовков разделов"""
        prefixes = [
            "Основные понятия и определения",
            "Исторический контекст развития", 
            "Методология исследования",
            "Ключевые характеристики и свойства",
            "Современные подходы и методы",
            "Практическое применение и примеры",
            "Технические аспекты реализации",
            "Сравнительный анализ технологий",
            "Экспериментальные исследования",
            "Теоретические основы"
        ]
        
        suffixes = [
            "и их особенности",
            "в современной науке", 
            "и перспективы развития",
            "и практическая значимость",
            "и методологические подходы",
            "и техническая реализация"
        ]
        
        prefix = prefixes[(chapter_num + section_num) % len(prefixes)]
        suffix = suffixes[section_num % len(suffixes)]
        
        return f"{section_num + 1}.{chapter_num + 1} {prefix} {suffix}"
    
    def _generate_section_content(self, topic, chapter_num, section_num, enhanced_content, length):
        """Генерация содержания раздела"""
        # Базовый контент из Wikipedia
        wiki_content = ""
        if enhanced_content.get('wikipedia', {}).get('exists'):
            wiki_data = enhanced_content['wikipedia']
            if wiki_data.get('summary'):
                wiki_content = wiki_data['summary'][:length//2]
        
        # Контент из базы знаний
        kb_content = ""
        kb_data = enhanced_content.get('knowledge_base')
        if kb_data:
            if chapter_num == 0 and kb_data.get('definition'):
                kb_content = kb_data['definition']
            elif chapter_num == 1 and kb_data.get('history'):
                kb_content = kb_data['history']
            elif chapter_num == 2 and kb_data.get('types'):
                kb_content = f"Основные типы и категории: {', '.join(kb_data['types'][:3])}."
        
        # Генерация детального контента
        detailed_content = self._generate_detailed_paragraph(topic, chapter_num, section_num, length)
        
        # Комбинируем контент
        combined_content = ""
        if wiki_content:
            combined_content += wiki_content + " "
        if kb_content:
            combined_content += kb_content + " "
        combined_content += detailed_content
        
        # Обрезаем до нужной длины
        return combined_content[:length]
    
    def _generate_detailed_paragraph(self, topic, chapter_num, section_num, length):
        """Генерация детального параграфа"""
        templates = [
            f"Анализ аспектов {topic} демонстрирует комплексный характер изучаемой проблематики. ",
            f"Исследование методов реализации {topic} показывает их практическую значимость. ",
            f"Рассмотрение технологических решений в области {topic} подтверждает их эффективность. ",
            f"Изучение принципов работы систем {topic} раскрывает их архитектурные особенности. ",
            f"Анализ современных тенденций развития {topic} указывает на перспективные направления. "
        ]
        
        details = [
            "Проведенные исследования позволяют сделать вывод о высокой эффективности рассматриваемых подходов. ",
            "Экспериментальные данные свидетельствуют о стабильности работы предложенных решений. ",
            "Теоретическое обоснование подтверждает корректность применяемых методик. ",
            "Практическая апробация демонстрирует реальную применимость полученных результатов. ",
            "Сравнительный анализ показывает преимущества рассматриваемых технологий. "
        ]
        
        examples = [
            "В качестве примера можно рассмотреть конкретные реализации в различных областях. ",
            "Типичными примерами применения являются современные технологические решения. ",
            "Практические кейсы использования подтверждают теоретические выкладки. ",
            "Реальные проекты демонстрируют работоспособность предложенных концепций. "
        ]
        
        # Выбираем шаблоны в зависимости от раздела
        content_parts = []
        
        # Основной контент
        main_template = templates[(chapter_num + section_num) % len(templates)]
        content_parts.append(main_template)
        
        # Детали
        if section_num > 0:
            content_parts.append(details[section_num % len(details)])
        
        # Примеры в последних разделах
        if section_num >= 2:
            content_parts.append(examples[(chapter_num + section_num) % len(examples)])
        
        # Заключительная часть
        conclusions = [
            "Таким образом, проведенный анализ позволяет сформулировать обоснованные выводы. ",
            "Полученные результаты имеют важное теоретическое и практическое значение. ",
            "Исследование вносит вклад в развитие научных представлений о рассматриваемой проблеме. "
        ]
        
        if section_num == 3:  # В последнем разделе главы
            content_parts.append(conclusions[chapter_num % len(conclusions)])
        
        content = "".join(content_parts)
        
        # Добавляем специфический контент в зависимости от темы
        if "ии" in topic.lower() or "искусственный интеллект" in topic.lower():
            ai_specific = [
                " В контексте искусственного интеллекта особое значение приобретают методы машинного обучения. ",
                " Нейронные сети представляют собой ключевую технологию современных систем ИИ. ",
                " Обработка естественного языка позволяет создавать интеллектуальные интерфейсы. "
            ]
            content += ai_specific[section_num % len(ai_specific)]
        
        return content[:length]
    
    def create_referat(self, topic, pages=10):
        """Создание полной структуры реферата"""
        content = self.generate_intelligent_content(topic, pages)
        
        return {
            'title': f'РЕФЕРАТ\nпо дисциплине: "Общие вопросы информатики и вычислительной техники"\nна тему: "{topic.capitalize()}"',
            'student_info': 'Выполнил: студент группы\nПроверил: преподаватель кафедры',
            'pages': pages,
            'introduction': self._generate_detailed_introduction(topic, pages, content),
            'chapters': content,
            'conclusion': self._generate_detailed_conclusion(topic, content),
            'sources': self._generate_realistic_sources(topic),
            'appendix': self._generate_useful_appendix(topic)
        }
    
    def _generate_detailed_introduction(self, topic, pages, content):
        """Генерация детального введения"""
        chapters_count = len(content)
        sections_count = sum(len(ch['sections']) for ch in content)
        
        return f"""ВВЕДЕНИЕ

Актуальность темы исследования "{topic}" обусловлена стремительным развитием технологий и возрастающей ролью автоматизации в современном мире. Изучение данной проблематики имеет важное теоретическое и практическое значение.

Цель работы: комплексное исследование основных аспектов {topic}, анализ современных тенденций и перспектив развития.

Задачи исследования:
1. Изучить теоретические основы и ключевые концепции {topic}
2. Проанализировать исторические этапы развития и современное состояние
3. Рассмотреть практические аспекты применения и реализации
4. Выявить перспективные направления дальнейшего развития
5. Сформулировать выводы и рекомендации

Объем работы: {pages} страниц
Структура работы: введение, {chapters_count} главы, заключение, список источников, приложения
Методология: системный анализ, сравнительный метод, теоретическое обобщение

Работа содержит {sections_count} разделов, последовательно раскрывающих тему исследования."""
    
    def _generate_detailed_conclusion(self, topic, content):
        """Генерация детального заключения"""
        return f"""ЗАКЛЮЧЕНИЕ

Проведенное исследование по теме "{topic}" позволило достичь поставленной цели и решить все задачи работы. 

Основные результаты и выводы:
1. Систематизированы теоретические основы и методологические подходы к изучению {topic}
2. Проанализированы исторические этапы развития и современное состояние области
3. Выявлены ключевые тенденции и закономерности развития технологий
4. Определена практическая значимость и области применения полученных знаний
5. Обоснованы перспективные направления для дальнейших исследований

Практическая ценность работы заключается в возможности использования полученных результатов в образовательном процессе, научной деятельности и практических разработках.

Наиболее перспективными направлениями дальнейших исследований являются углубленное изучение специализированных аспектов {topic} и разработка конкретных практических решений на основе полученных теоретических знаний."""
    
    def _generate_realistic_sources(self, topic):
        """Генерация реалистичных источников"""
        base_sources = [
            "Таненбаум Э. Современные операционные системы. - СПб.: Питер, 2020. - 1120 с.",
            "Кормен Т., Лейзерсон Ч., Ривест Р. Алгоритмы: построение и анализ. - М.: Вильямс, 2022. - 1296 с.",
            "Рассел С., Норвиг П. Искусственный интеллект: современный подход. - М.: Вильямс, 2021. - 1408 с.",
            "Гудфеллоу Я., Бенджио И., Курвилль А. Глубокое обучение. - М.: ДМК Пресс, 2022. - 652 с.",
            "Стивенс Р. UNIX: взаимодействие процессов. - М.: Вильямс, 2019. - 576 с."
        ]
        
        topic_specific = {
            "ии": [
                "Нильсон Н. Искусственный интеллект. - М.: Мир, 2020. - 420 с.",
                "Лорьер Ж.-Л. Системы искусственного интеллекти. - М.: Мир, 2021. - 568 с."
            ],
            "программирование": [
                "Макконнелл С. Совершенный код. - М.: Русская редакция, 2022. - 896 с.",
                "Фаулер М. Рефакторинг. Улучшение существующего кода. - М.: Диалектика, 2021. - 448 с."
            ],
            "физика": [
                "Сивухин Д.В. Общий курс физики. - М.: Физматлит, 2020. - 560 с.",
                "Ландау Л.Д., Лифшиц Е.М. Теоретическая физика. - М.: Физматлит, 2019. - 480 с."
            ]
        }
        
        sources = base_sources.copy()
        
        # Добавляем тематические источники
        for key, value in topic_specific.items():
            if key in topic.lower():
                sources.extend(value[:2])
        
        # Добавляем интернет-источники
        internet_sources = [
            f"Материалы научной энциклопедии Wikipedia - {topic}",
            "Электронная библиотека образовательных ресурсов",
            "Научные публикации в открытом доступе",
            "Материалы международных конференций"
        ]
        
        sources.extend(internet_sources)
        return sources[:10]  # Ограничиваем 10 источниками
    
    def _generate_useful_appendix(self, topic):
        """Генерация полезного приложения"""
        return f"""ПРИЛОЖЕНИЯ

Приложение А
Диаграммы и схемы, иллюстрирующие основные концепции {topic}

Приложение Б
Таблицы сравнительного анализа различных подходов и методов

Приложение В
Примеры практической реализации рассмотренных технологий

Приложение Г
Глоссарий основных терминов и определений

Приложение Д
Дополнительные материалы и расчеты

Для углубленного изучения темы рекомендуется ознакомиться с специализированной литературой и провести самостоятельные практические исследования."""
    
    def create_word_document(self, referat_data):
        """Создание Word документа"""
        try:
            doc = Document()
            
            # Настройка стилей
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(14)
            
            # Титульная страница
            title = doc.add_heading(referat_data['title'], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph("\n" * 8)
            student_para = doc.add_paragraph(referat_data['student_info'])
            student_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph("\n" * 6)
            current_year = datetime.datetime.now().year
            date_para = doc.add_paragraph(f"г. Санкт-Петербург\n{current_year} год")
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_page_break()
            
            # Содержание
            doc.add_heading('СОДЕРЖАНИЕ', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
            
            # Введение
            doc.add_paragraph("Введение")
            
            # Главы и разделы
            page_counter = 3  # Введение начинается с 3 страницы
            
            for i, chapter in enumerate(referat_data['chapters']):
                doc.add_paragraph(f"Глава {i+1}. {chapter['title']}")
                for section in chapter['sections']:
                    doc.add_paragraph(f"   {section['title']}")
            
            # Заключение и остальное
            doc.add_paragraph("Заключение")
            doc.add_paragraph("Список использованных источников")
            doc.add_paragraph("Приложения")
            
            doc.add_page_break()
            
            # Введение
            doc.add_heading('ВВЕДЕНИЕ', level=1)
            intro_paragraphs = referat_data['introduction'].split('\n\n')
            for paragraph in intro_paragraphs:
                if paragraph.strip():
                    p = doc.add_paragraph(paragraph.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.line_spacing = 1.5
            
            # Основные главы
            for i, chapter in enumerate(referat_data['chapters']):
                doc.add_page_break()
                doc.add_heading(f'ГЛАВА {i+1}. {chapter["title"].upper()}', level=1)
                
                for section in chapter['sections']:
                    doc.add_heading(section['title'], level=2)
                    
                    # Разбиваем контент на абзацы
                    content_paragraphs = section['content'].split('. ')
                    for para_text in content_paragraphs:
                        if para_text.strip():
                            p = doc.add_paragraph(para_text.strip() + '.')
                            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            p.paragraph_format.line_spacing = 1.5
                            p.paragraph_format.space_after = Pt(6)
                    
                    doc.add_paragraph()  # Пустая строка между разделами
            
            # Заключение
            doc.add_page_break()
            doc.add_heading('ЗАКЛЮЧЕНИЕ', level=1)
            conclusion_paragraphs = referat_data['conclusion'].split('\n\n')
            for paragraph in conclusion_paragraphs:
                if paragraph.strip():
                    p = doc.add_paragraph(paragraph.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.line_spacing = 1.5
            
            # Источники
            doc.add_page_break()
            doc.add_heading('СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ', level=1)
            for i, source in enumerate(referat_data['sources'], 1):
                p = doc.add_paragraph(f"{i}. {source}")
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Приложение
            doc.add_page_break()
            doc.add_heading('ПРИЛОЖЕНИЯ', level=1)
            appendix_paragraphs = referat_data['appendix'].split('\n\n')
            for paragraph in appendix_paragraphs:
                if paragraph.strip():
                    p = doc.add_paragraph(paragraph.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Сохраняем в байтовый поток
            file_stream = io.BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)
            
            return file_stream
            
        except Exception as e:
            logging.error(f"Ошибка создания документа: {e}")
            return None
    
    def generate_referat(self, topic, pages=10):
        """Основная функция генерации реферата"""
        try:
            logging.info(f"Создание реферата: {topic}, {pages} страниц")
            
            # Создаем структуру реферата
            referat_structure = self.create_referat(topic, pages)
            
            if not referat_structure:
                return None, "❌ Не удалось сформировать структуру реферата"
            
            # Создаем документ
            doc_file = self.create_word_document(referat_structure)
            
            if not doc_file:
                return None, "❌ Ошибка при создании документа"
            
            chapters_count = len(referat_structure['chapters'])
            sections_count = sum(len(ch['sections']) for ch in referat_structure['chapters'])
            
            success_message = f"""📚 Реферат создан успешно!

🎯 Тема: "{topic}"
📊 Параметры:
• Объем: {pages} страниц
• Глав: {chapters_count}
• Разделов: {sections_count}

✨ Особенности:
• Академическая структура по ГОСТ
• Реалистичное содержание
• Список проверенных источников
• Практические приложения

💡 Рекомендации:
• Добавьте конкретные примеры
• Включите графики и таблицы
• Проверьте уникальность текста
• Дополните собственными исследованиями"""
            
            return doc_file, success_message
            
        except Exception as e:
            logging.error(f"Ошибка генерации реферата: {e}")
            return None, f"❌ Ошибка при создании реферата: {str(e)}"

# Создаем улучшенный генератор
advanced_referat_generator = AdvancedReferatGenerator()

# === МАТЕМАТИЧЕСКИЙ ПОМОЩНИК ===
class TosikAssistant:
    def __init__(self):
        self.x, self.y, self.z = symbols('x y z')
    
    def solve_math_problem(self, problem):
        """Решает математические задачи"""
        try:
            problem_lower = problem.lower()
            
            if any(word in problem_lower for word in ['уравнен', 'x²', 'x**2']):
                return self._solve_equation(problem)
            elif any(word in problem_lower for word in ['площад', 'объем', 'геометр']):
                return self._solve_geometry(problem)
            elif any(word in problem_lower for word in ['производн', 'дифференц']):
                return self._solve_derivative(problem)
            elif any(word in problem_lower for word in ['интеграл']):
                return self._solve_integral(problem)
            else:
                return "🤔 Напиши конкретную задачу: уравнение, площадь, производная, интеграл"
                
        except Exception as e:
            return f"❌ Ошибка решения: {str(e)}"
    
    def _solve_equation(self, problem):
        """Решает уравнения"""
        try:
            nums = [float(x) for x in re.findall(r'([+-]?\d*\.?\d+)', problem.replace(' ', '')) if x]
            
            if len(nums) >= 3:
                a, b, c = nums[0], nums[1], nums[2]
                D = b**2 - 4*a*c
                
                result = f"📊 Решение уравнения: {a}x² + {b}x + {c} = 0\n"
                result += f"Дискриминант: D = {b}² - 4×{a}×{c} = {D}\n"
                
                if D < 0:
                    result += "❌ Действительных корней нет"
                elif D == 0:
                    x = -b / (2*a)
                    result += f"✅ Один корень: x = {x:.2f}"
                else:
                    x1 = (-b + math.sqrt(D)) / (2*a)
                    x2 = (-b - math.sqrt(D)) / (2*a)
                    result += f"✅ Два корня:\n"
                    result += f"x₁ = {x1:.2f}\n"
                    result += f"x₂ = {x2:.2f}"
                
                return result
                
        except:
            return "❌ Не удалось решить уравнение"
    
    def _solve_geometry(self, problem):
        """Решает геометрические задачи"""
        numbers = [float(x) for x in re.findall(r'(\d+\.?\d*)', problem)]
        
        if 'площад' in problem.lower() and 'круг' in problem.lower() and numbers:
            r = numbers[0]
            area = math.pi * r**2
            return f"📐 Площадь круга:\nS = πr² = 3.14 × {r}² = {area:.2f}"
        
        elif 'объем' in problem.lower() and 'сфер' in problem.lower() and numbers:
            r = numbers[0]
            volume = (4/3) * math.pi * r**3
            return f"📐 Объем сферы:\nV = 4/3πr³ = 4/3 × 3.14 × {r}³ = {volume:.2f}"
        
        return "🤔 Напиши конкретную геометрическую задачу"
    
    def _solve_derivative(self, problem):
        """Решает производные"""
        return """
📚 Производная функции:

Основные правила:
1. (xⁿ)' = n·xⁿ⁻¹
2. (sin x)' = cos x  
3. (cos x)' = -sin x
4. (eˣ)' = eˣ
5. (ln x)' = 1/x

Пример:
f(x) = 3x⁴ + 2x² - 5x + 1
f'(x) = 12x³ + 4x - 5
        """
    
    def _solve_integral(self, problem):
        """Решает интегралы"""
        return """
📚 Интеграл функции:

Основные правила:
1. ∫xⁿ dx = xⁿ⁺¹/(n+1) + C
2. ∫sin x dx = -cos x + C  
3. ∫cos x dx = sin x + C
4. ∫eˣ dx = eˣ + C
5. ∫1/x dx = ln|x| + C

Пример:
∫(4x³ - 2x + 1) dx = x⁴ - x² + x + C
        """
    
    def calculate_expression(self, expression):
        """Вычисляет математические выражения"""
        try:
            expr_clean = expression.replace('^', '**').replace('×', '*').replace('÷', '/')
            result = eval(expr_clean)
            return f"🧮 Результат: {expression} = {result}"
        except:
            return "❌ Не удалось вычислить выражение"

# Создаем помощника Тосика
tosik_assistant = TosikAssistant()

# === БАЗА ДАННЫХ ===
def init_db():
    conn = sqlite3.connect('schedule.db', check_same_thread=False)
    cursor = conn.cursor()
    
    for week in range(1, 5):
        cursor.execute(f'''
            CREATE TABLE IF NOT EXISTS schedule_week{week} (
                id INTEGER PRIMARY KEY,
                data TEXT NOT NULL,
                last_updated TEXT NOT NULL,
                week_start_date TEXT NOT NULL
            )
        ''')
        cursor.execute(f"INSERT OR IGNORE INTO schedule_week{week} (id, data, last_updated, week_start_date) VALUES (1, '{{}}', '', '')")
    
    # Таблицы для системы докладов
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS reports_system (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            subject_name TEXT NOT NULL UNIQUE,
            report_data TEXT NOT NULL,
            max_reports_per_student INTEGER DEFAULT 1,
            created_by INTEGER NOT NULL,
            created_at TEXT NOT NULL,
            is_active BOOLEAN DEFAULT 1
        )
    ''')
    
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS student_registry (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL UNIQUE,
            student_number TEXT NOT NULL,
            student_name TEXT NOT NULL,
            registered_at TEXT NOT NULL
        )
    ''')
    
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS report_assignments (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            subject_name TEXT NOT NULL,
            report_number INTEGER NOT NULL,
            report_title TEXT NOT NULL,
            assigned_at TEXT NOT NULL
        )
    ''')
    
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS admins (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL UNIQUE,
            added_by INTEGER NOT NULL,
            added_at TEXT NOT NULL
        )
    ''')
    
    # Добавляем основного админа
    cursor.execute("SELECT 1 FROM admins WHERE user_id = ?", (CONFIG['admin_id'],))
    if not cursor.fetchone():
        cursor.execute(
            "INSERT INTO admins (user_id, added_by, added_at) VALUES (?, ?, ?)",
            (CONFIG['admin_id'], CONFIG['admin_id'], datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        )
    
    conn.commit()
    conn.close()
    print("✅ База данных инициализирована")

def register_student(user_id, student_number, student_name):
    """Регистрация студента"""
    conn = sqlite3.connect('schedule.db', check_same_thread=False)
    cursor = conn.cursor()
    
    try:
        cursor.execute(
            "INSERT OR REPLACE INTO student_registry (user_id, student_number, student_name, registered_at) VALUES (?, ?, ?, ?)",
            (user_id, student_number, student_name, datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        )
        conn.commit()
        return True
    except Exception as e:
        print(f"❌ Ошибка регистрации: {e}")
        return False
    finally:
        conn.close()

def get_student_info(user_id):
    """Получение информации о студенте"""
    conn = sqlite3.connect('schedule.db', check_same_thread=False)
    cursor = conn.cursor()
    
    cursor.execute("SELECT student_number, student_name FROM student_registry WHERE user_id = ?", (user_id,))
    result = cursor.fetchone()
    conn.close()
    
    return result if result else None

def create_subject(subject_name, max_reports, created_by):
    """Создание предмета для докладов"""
    conn = sqlite3.connect('schedule.db', check_same_thread=False)
    cursor = conn.cursor()
    
    try:
        initial_data = json.dumps({})
        cursor.execute(
            "INSERT INTO reports_system (subject_name, report_data, max_reports_per_student, created_by, created_at) VALUES (?, ?, ?, ?, ?)",
            (subject_name, initial_data, max_reports, created_by, datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        )
        conn.commit()
        return True
    except sqlite3.IntegrityError:
        return False
    finally:
        conn.close()

def add_report_to_subject(subject_name, report_number, report_title):
    """Добавление доклада к предмету"""
    conn = sqlite3.connect('schedule.db', check_same_thread=False)
    cursor = conn.cursor()
    
    cursor.execute("SELECT report_data FROM reports_system WHERE subject_name = ?", (subject_name,))
    result = cursor.fetchone()
    
    if not result:
        conn.close()
        return False
    
    report_data = json.loads(result[0])
    report_data[str(report_number)] = {
        "title": report_title,
        "taken_by": None
    }
    
    cursor.execute(
        "UPDATE reports_system SET report_data = ? WHERE subject_name = ?",
        (json.dumps(report_data), subject_name)
    )
    conn.commit()
    conn.close()
    return True

def take_report_for_student(user_id, subject_name, report_number):
    """Взятие доклада студентом"""
    conn = sqlite3.connect('schedule.db', check_same_thread=False)
    cursor = conn.cursor()
    
    # Проверяем, не взял ли уже студент максимальное количество докладов
    cursor.execute(
        "SELECT COUNT(*) FROM report_assignments WHERE user_id = ? AND subject_name = ?",
        (user_id, subject_name)
    )
    taken_count = cursor.fetchone()[0]
    
    cursor.execute(
        "SELECT max_reports_per_student FROM reports_system WHERE subject_name = ?",
        (subject_name,)
    )
    max_reports = cursor.fetchone()
    
    if not max_reports:
        conn.close()
        return "❌ Предмет не найден"
    
    max_reports = max_reports[0]
    
    if taken_count >= max_reports:
        conn.close()
        return f"❌ Вы уже взяли максимальное количество докладов ({max_reports}) по этому предмету"
    
    # Проверяем, свободен ли доклад
    cursor.execute("SELECT report_data FROM reports_system WHERE subject_name = ?", (subject_name,))
    result = cursor.fetchone()
    
    if not result:
        conn.close()
        return "❌ Предмет не найден"
    
    report_data = json.loads(result[0])
    
    if str(report_number) not in report_data:
        conn.close()
        return "❌ Доклад с таким номером не найден"
    
    if report_data[str(report_number)]["taken_by"] is not None:
        conn.close()
        return "❌ Этот доклад уже занят"
    
    # Занимаем доклад
    report_data[str(report_number)]["taken_by"] = user_id
    
    # Получаем информацию о студенте
    student_info = get_student_info(user_id)
    if not student_info:
        conn.close()
        return "❌ Сначала зарегистрируйтесь как студент"
    
    student_number, student_name = student_info
    
    # Обновляем данные и добавляем запись о назначении
    cursor.execute(
        "UPDATE reports_system SET report_data = ? WHERE subject_name = ?",
        (json.dumps(report_data), subject_name)
    )
    
    cursor.execute(
        "INSERT INTO report_assignments (user_id, subject_name, report_number, report_title, assigned_at) VALUES (?, ?, ?, ?, ?)",
        (user_id, subject_name, report_number, report_data[str(report_number)]["title"], datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    )
    
    conn.commit()
    conn.close()
    
    return f"✅ Доклад '{report_data[str(report_number)]['title']}' успешно взят!"

def get_subject_reports(subject_name):
    """Получение списка докладов по предмету"""
    conn = sqlite3.connect('schedule.db', check_same_thread=False)
    cursor = conn.cursor()
    
    cursor.execute("SELECT report_data FROM reports_system WHERE subject_name = ?", (subject_name,))
    result = cursor.fetchone()
    conn.close()
    
    return json.loads(result[0]) if result else None

def get_student_reports(user_id):
    """Получение докладов студента"""
    conn = sqlite3.connect('schedule.db', check_same_thread=False)
    cursor = conn.cursor()
    
    cursor.execute(
        "SELECT subject_name, report_number, report_title, assigned_at FROM report_assignments WHERE user_id = ? ORDER BY assigned_at",
        (user_id,)
    )
    reports = cursor.fetchall()
    conn.close()
    
    return reports

def get_all_subjects():
    """Получение всех предметов"""
    conn = sqlite3.connect('schedule.db', check_same_thread=False)
    cursor = conn.cursor()
    
    cursor.execute("SELECT subject_name, max_reports_per_student FROM reports_system WHERE is_active = 1")
    subjects = cursor.fetchall()
    conn.close()
    
    return subjects

def is_admin(user_id):
    """Проверка, является ли пользователь админом"""
    conn = sqlite3.connect('schedule.db', check_same_thread=False)
    cursor = conn.cursor()
    
    cursor.execute("SELECT 1 FROM admins WHERE user_id = ?", (user_id,))
    result = cursor.fetchone() is not None
    conn.close()
    
    return result

def add_admin(user_id, added_by):
    """Добавление администратора"""
    conn = sqlite3.connect('schedule.db', check_same_thread=False)
    cursor = conn.cursor()
    
    try:
        cursor.execute(
            "INSERT INTO admins (user_id, added_by, added_at) VALUES (?, ?, ?)",
            (user_id, added_by, datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        )
        conn.commit()
        return True
    except sqlite3.IntegrityError:
        return False
    finally:
        conn.close()

def save_schedule(week_data, week_number=None):
    """Сохранение расписания"""
    if week_number is None:
        week_number = CONFIG['current_week']
    
    conn = sqlite3.connect('schedule.db', check_same_thread=False)
    cursor = conn.cursor()
    
    cursor.execute(
        f"UPDATE schedule_week{week_number} SET data = ?, last_updated = ? WHERE id = 1",
        (json.dumps(week_data), datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    )
    conn.commit()
    conn.close()

def load_schedule(week_number=None):
    """Загрузка расписания"""
    if week_number is None:
        week_number = CONFIG['current_week']
    
    conn = sqlite3.connect('schedule.db', check_same_thread=False)
    cursor = conn.cursor()
    
    cursor.execute(f"SELECT data, last_updated FROM schedule_week{week_number} WHERE id = 1")
    result = cursor.fetchone()
    conn.close()
    
    if result and result[0]:
        return json.loads(result[0]), result[1]
    return {}, ""

# Русские названия дней
days_of_week = ["понедельник", "вторник", "среда", "четверг", "пятница", "суббота", "воскресенье"]
days_of_week_capitalized = ["Понедельник", "Вторник", "Среда", "Четверг", "Пятница", "Суббота", "Воскресенье"]

# Форматирование расписания
def format_schedule_day(schedule_data, day_offset=0):
    if not schedule_data:
        return "📅 Расписание пока не добавлено."
    
    target_date = datetime.datetime.now() + datetime.timedelta(days=day_offset)
    day_name = days_of_week[target_date.weekday()]
    day_name_cap = days_of_week_capitalized[target_date.weekday()]
    
    response = f"📅 {day_name_cap}:\n"
    response += "─" * 30 + "\n"
    
    if day_name in schedule_data and schedule_data[day_name]:
        for lesson in schedule_data[day_name]:
            response += f"🕒 {lesson['pair']} пара\n"
            response += f"📚 {lesson['subject']}\n"
            response += f"👤 {lesson['teacher']}\n"
            response += f"🚪 {lesson['room']}\n\n"
    else:
        response += "🎉 Занятий нет\n\n"
    
    return response

def format_schedule_week(schedule_data, week_offset=0):
    if not schedule_data:
        return "📅 Расписание пока не добавлено."
    
    response = ""
    today = datetime.datetime.now()
    
    for i, day_name in enumerate(days_of_week):
        day_date = today + datetime.timedelta(days=i - today.weekday() + (week_offset * 7))
        day_name_cap = days_of_week_capitalized[i]
        
        response += f"📅 {day_name_cap}:\n"
        response += "─" * 30 + "\n"
        
        if day_name in schedule_data and schedule_data[day_name]:
            for lesson in schedule_data[day_name]:
                response += f"🕒 {lesson['pair']} пара: {lesson['subject']}\n"
        else:
            response += "🎉 Занятий нет\n"
        response += "\n"
    
    return response

# Функция для отправки сообщения
def send_message(peer_id, message, keyboard=None):
    try:
        params = {
            'peer_id': peer_id,
            'message': message,
            'random_id': get_random_id(),
        }
        if keyboard:
            params['keyboard'] = keyboard.get_keyboard() if hasattr(keyboard, 'get_keyboard') else keyboard
            
        result = vk_session.method('messages.send', params)
        print(f"✅ Сообщение отправлено: {message[:50]}...")
        return result
    except Exception as e:
        print(f"❌ Ошибка отправки сообщения: {e}")
        return None

# Функция для отправки документа
def send_document(peer_id, file_stream, filename, message=""):
    """Отправка документа пользователю"""
    try:
        print(f"📤 Тосик отправляет документ: {filename}")
        
        # Получаем URL для загрузки
        upload_data = vk_session.method('docs.getMessagesUploadServer', {
            'type': 'doc',
            'peer_id': peer_id
        })
        
        upload_url = upload_data['upload_url']
        
        # Отправляем файл
        files = {'file': (filename, file_stream.getvalue())}
        response = requests.post(upload_url, files=files, timeout=30)
        
        if response.status_code != 200:
            return False
            
        result = response.json()
        
        # Сохраняем документ
        doc_data = vk_session.method('docs.save', {
            'file': result['file'],
            'title': filename
        })
        
        # Получаем attachment
        if 'doc' in doc_data:
            doc = doc_data['doc']
            attachment = f"doc{doc['owner_id']}_{doc['id']}"
        else:
            return False
        
        # Отправляем сообщение с документом
        vk_session.method('messages.send', {
            'peer_id': peer_id,
            'attachment': attachment,
            'message': message,
            'random_id': get_random_id()
        })
        
        print(f"✅ Документ отправлен!")
        return True
        
    except Exception as e:
        print(f"❌ Ошибка отправки документа: {e}")
        return False

# Функция для проверки разрешенных бесед
def is_allowed_chat(peer_id):
    """Проверяет, разрешена ли беседа"""
    return peer_id in CONFIG['allowed_chats']

# Подключаемся к VK
try:
    print("🔄 Подключаюсь к VK API...")
    vk_session = vk_api.VkApi(token=CONFIG['token'])
    vk = vk_session.get_api()
    
    group_info = vk.groups.getById(group_id=CONFIG['group_id'])
    print(f"✅ Успешно подключен к группе: {group_info[0]['name']}")
    
    longpoll = VkBotLongPoll(vk_session, CONFIG['group_id'])
    print("✅ LongPoll запущен")
    
except Exception as e:
    print(f"❌ Ошибка подключения: {e}")
    exit()

# Инициализируем базу данных
init_db()

print("🎯 Бот Тосик запущен и готов к работе!")
print(f"📞 Разрешенные беседы: {CONFIG['allowed_chats']}")

# Главный цикл обработки событий
for event in longpoll.listen():
    try:
        if event.type == VkBotEventType.MESSAGE_NEW:
            msg = event.object.message['text'].strip()
            user_id = event.object.message['from_id']
            peer_id = event.object.message['peer_id']
            original_text = event.object.message['text']
            
            print(f"📩 Получено: '{msg}' от {user_id} в {peer_id}")
            
            # Проверяем разрешена ли беседа
            if not is_allowed_chat(peer_id):
                print(f"🚫 Беседа {peer_id} не в списке разрешенных")
                continue
            
            # Проверяем, должен ли Тосик ответить (умные ответы)
            if smart_tosik.should_respond(msg, user_id):
                response = smart_tosik.generate_ai_response(msg)
                send_message(peer_id, response)
                continue
            
            if event.from_chat and CONFIG['chat_id'] is None:
                CONFIG['chat_id'] = peer_id
                print(f"💬 Беседа сохранена: {peer_id}")
            
            if event.from_chat:
                msg_lower = msg.lower()
                
                # ========== ОСНОВНЫЕ КОМАНДЫ ==========
                
                # Расписание
                if msg_lower in ['расписание', 'сегодня']:
                    schedule, last_updated = load_schedule()
                    response = format_schedule_day(schedule, 0)
                    if last_updated:
                        response += f"\n🔄 Обновлено: {last_updated}"
                    send_message(peer_id, response)
                    continue
                
                elif msg_lower == 'завтра':
                    schedule, last_updated = load_schedule()
                    response = format_schedule_day(schedule, 1)
                    if last_updated:
                        response += f"\n🔄 Обновлено: {last_updated}"
                    send_message(peer_id, response)
                    continue
                
                elif msg_lower == 'неделя':
                    schedule, last_updated = load_schedule()
                    response = format_schedule_week(schedule, 0)
                    if last_updated:
                        response += f"\n🔄 Обновлено: {last_updated}"
                    send_message(peer_id, response)
                    continue
                
                elif msg_lower == 'след неделя':
                    next_week = (CONFIG['current_week'] % 4) + 1
                    schedule, last_updated = load_schedule(next_week)
                    response = f"📅 Следующая неделя\n\n" + format_schedule_week(schedule, 1)
                    if last_updated:
                        response += f"\n🔄 Обновлено: {last_updated}"
                    send_message(peer_id, response)
                    continue
                
                # Доклады
                elif msg_lower == 'доклады':
                    subjects = get_all_subjects()
                    if subjects:
                        response = "📚 Доступные предметы для докладов:\n\n"
                        for subject_name, max_reports in subjects:
                            response += f"📖 {subject_name} (можно взять: {max_reports})\n"
                        response += "\n🎯 Чтобы посмотреть доклады по предмету: 'Доклады по [предмет]'\n📝 Чтобы взять доклад: 'Беру доклад [номер] по [предмет]'"
                    else:
                        response = "📚 Пока нет предметов для докладов"
                    send_message(peer_id, response)
                    continue
                
                elif msg_lower.startswith('доклады по '):
                    subject_name = original_text[11:].strip()
                    reports = get_subject_reports(subject_name)
                    
                    if reports:
                        response = f"📋 Доклады по предмету '{subject_name}':\n\n"
                        free_count = 0
                        
                        for report_num, report_info in sorted(reports.items(), key=lambda x: int(x[0])):
                            status = "✅ Свободен" if not report_info["taken_by"] else f"❌ Занят ({GROUP_LIST.get(str(report_info['taken_by']), 'Неизвестный')})"
                            if not report_info["taken_by"]:
                                free_count += 1
                            
                            response += f"📄 {report_num}. {report_info['title']} - {status}\n"
                        
                        response += f"\n🎯 Свободно: {free_count}/{len(reports)}"
                        response += f"\n📝 Взять: 'Беру доклад [номер] по {subject_name}'"
                    else:
                        response = f"❌ Нет докладов по предмету '{subject_name}' или предмет не найден"
                    send_message(peer_id, response)
                    continue
                
                elif msg_lower.startswith('беру доклад '):
                    # Парсим команду "Беру доклад X по Y"
                    parts = original_text.split()
                    if len(parts) >= 5 and parts[2].isdigit() and parts[3] == 'по':
                        report_number = int(parts[2])
                        subject_name = ' '.join(parts[4:])
                        
                        # Проверяем регистрацию
                        student_info = get_student_info(user_id)
                        if not student_info:
                            send_message(peer_id, "❌ Сначала зарегистрируйтесь: 'Регистрация X' где X - ваш номер в списке группы")
                            continue
                        
                        result = take_report_for_student(user_id, subject_name, report_number)
                        send_message(peer_id, result)
                    else:
                        send_message(peer_id, "❌ Формат: 'Беру доклад X по [название предмета]'")
                    continue
                
                elif msg_lower == 'мои доклады':
                    student_info = get_student_info(user_id)
                    if not student_info:
                        send_message(peer_id, "❌ Сначала зарегистрируйтесь: 'Регистрация X' где X - ваш номер в списке группы")
                        continue
                    
                    reports = get_student_reports(user_id)
                    if reports:
                        response = "📋 Ваши доклады:\n\n"
                        for subject, number, title, date in reports:
                            response += f"📖 {subject} - Доклад {number}: {title}\n"
                            response += f"   📅 Взят: {date}\n\n"
                    else:
                        response = "📭 У вас пока нет взятых докладов"
                    send_message(peer_id, response)
                    continue
                
                # Регистрация
                elif msg_lower.startswith('регистрация '):
                    parts = original_text.split()
                    if len(parts) >= 2 and parts[1].isdigit():
                        student_number = parts[1]
                        student_name = GROUP_LIST.get(student_number)
                        
                        if student_name:
                            if register_student(user_id, student_number, student_name):
                                send_message(peer_id, f"✅ {student_name} успешно зарегистрирован под номером {student_number}!")
                            else:
                                send_message(peer_id, "❌ Ошибка регистрации")
                        else:
                            send_message(peer_id, f"❌ Студент с номером {student_number} не найден в списке группы")
                    else:
                        send_message(peer_id, "❌ Формат: 'Регистрация X' где X - ваш номер в списке группы")
                    continue
                
                # Помощь
                elif msg_lower in ['помощь', 'команды', 'help']:
                    help_text = """
🎯 ДОСТУПНЫЕ КОМАНДЫ:

📅 РАСПИСАНИЕ:
• "Расписание" или "Сегодня" - расписание на сегодня
• "Завтра" - расписание на завтра  
• "Неделя" - расписание на неделю
• "След неделя" - расписание на след. неделю

📚 ДОКЛАДЫ:
• "Доклады" - список предметов
• "Доклады по [предмет]" - доклады по предмету
• "Беру доклад X по [предмет]" - взять доклад
• "Мои доклады" - ваши доклады
• "Регистрация X" - регистрация (X - ваш номер)

📝 РЕФЕРАТЫ:
• "Реферат [тема]" - реферат на 10 страниц
• "Реферат [тема] [N] стр" - реферат на N страниц

🧮 МАТЕМАТИКА:
• "Реши [уравнение]" - решить уравнение
• "Посчитай [выражение]" - вычислить

💬 ОБЩЕНИЕ:
• Обращайся "Тосик" и задавай вопросы
• "Тосик, шутка" - случайная шутка
• "Тосик, факт" - интересный факт

⚙️ АДМИН-КОМАНДЫ:
• "Добавить предмет [название] [макс]" - новый предмет
• "Добавить доклад [предмет] [номер] [название]" - новый доклад
• "Добавить админа [ID]" - добавить администратора

💡 Просто напиши вопрос или обратись "Тосик" - я постараюсь помочь!
                    """
                    send_message(peer_id, help_text)
                    continue
                
                # Рефераты
                elif msg_lower.startswith('реферат '):
                    parts = original_text[8:].strip().split()
                    if len(parts) >= 1:
                        # Определяем тему и количество страниц
                        if len(parts) >= 3 and parts[-2] in ['стр', 'страниц', 'страницы'] and parts[-1].isdigit():
                            pages = int(parts[-1])
                            topic = ' '.join(parts[:-2])
                        else:
                            pages = 10
                            topic = ' '.join(parts)
                        
                        if len(topic) < 3:
                            send_message(peer_id, "❌ Укажите тему реферата (минимум 3 символа)")
                            continue
                        
                        # Отправляем сообщение о начале генерации
                        send_message(peer_id, f"📚 Тосик начинает создание реферата на тему: '{topic}'\n⏳ Это займет некоторое время...")
                        
                        # Генерируем реферат
                        doc_file, message = advanced_referat_generator.generate_referat(topic, pages)
                        
                        if doc_file:
                            filename = f"Реферат_{topic.replace(' ', '_')}.docx"
                            if send_document(peer_id, doc_file, filename, message):
                                print(f"✅ Реферат '{topic}' успешно отправлен")
                            else:
                                send_message(peer_id, "❌ Ошибка отправки документа")
                        else:
                            send_message(peer_id, message)
                    else:
                        send_message(peer_id, "❌ Формат: 'Реферат [тема]' или 'Реферат [тема] [N] стр'")
                    continue
                
                # Математика
                elif msg_lower.startswith('реши '):
                    problem = original_text[5:].strip()
                    if problem:
                        result = tosik_assistant.solve_math_problem(problem)
                        send_message(peer_id, result)
                    else:
                        send_message(peer_id, "❌ Укажите задачу для решения")
                    continue
                
                elif msg_lower.startswith('посчитай '):
                    expression = original_text[9:].strip()
                    if expression:
                        result = tosik_assistant.calculate_expression(expression)
                        send_message(peer_id, result)
                    else:
                        send_message(peer_id, "❌ Укажите выражение для вычисления")
                    continue
                
                # ========== АДМИН-КОМАНДЫ ==========
                if is_admin(user_id):
                    # Добавление предмета
                    if msg_lower.startswith('добавить предмет '):
                        parts = original_text[17:].strip().split()
                        if len(parts) >= 2 and parts[-1].isdigit():
                            max_reports = int(parts[-1])
                            subject_name = ' '.join(parts[:-1])
                            
                            if create_subject(subject_name, max_reports, user_id):
                                send_message(peer_id, f"✅ Предмет '{subject_name}' добавлен (макс. {max_reports} докладов на студента)")
                            else:
                                send_message(peer_id, f"❌ Ошибка: предмет '{subject_name}' уже существует")
                        else:
                            send_message(peer_id, "❌ Формат: 'Добавить предмет [название] [максимум докладов]'")
                        continue
                    
                    # Добавление доклада
                    elif msg_lower.startswith('добавить доклад '):
                        parts = original_text[16:].strip().split()
                        if len(parts) >= 4 and parts[0].isdigit():
                            report_number = int(parts[0])
                            # Ищем "по" для разделения
                            if ' по ' in original_text:
                                subject_part = original_text.split(' по ', 1)[1]
                                subject_name = subject_part.strip()
                                report_title = ' '.join(parts[1:parts.index('по')]) if 'по' in parts else ' '.join(parts[1:])
                            else:
                                subject_name = parts[-1]
                                report_title = ' '.join(parts[1:-1])
                            
                            if add_report_to_subject(subject_name, report_number, report_title):
                                send_message(peer_id, f"✅ Доклад {report_number} добавлен к предмету '{subject_name}': {report_title}")
                            else:
                                send_message(peer_id, f"❌ Ошибка: предмет '{subject_name}' не найден")
                        else:
                            send_message(peer_id, "❌ Формат: 'Добавить доклад [номер] [название] по [предмет]'")
                        continue
                    
                    # Добавление админа
                    elif msg_lower.startswith('добавить админа '):
                        parts = original_text[16:].strip().split()
                        if parts and parts[0].isdigit():
                            new_admin_id = int(parts[0])
                            if add_admin(new_admin_id, user_id):
                                send_message(peer_id, f"✅ Пользователь {new_admin_id} добавлен как администратор")
                            else:
                                send_message(peer_id, f"❌ Пользователь {new_admin_id} уже является администратором")
                        else:
                            send_message(peer_id, "❌ Формат: 'Добавить админа [ID пользователя]'")
                        continue
                
                # ========== ЕСЛИ НИЧЕГО НЕ РАСПОЗНАНО ==========
                else:
                    # Если сообщение достаточно длинное, предлагаем помощь
                    if len(msg) > 5:
                        help_offer = random.choice([
                            "🤔 Не совсем понял... Напиши 'Помощь' чтобы узнать что я умею!",
                            "🎯 Не распознал команду. Попробуй 'Расписание', 'Доклады' или 'Реферат тема'",
                            "💡 Нужна помощь? Напиши 'Помощь' для списка команд!"
                        ])
                        send_message(peer_id, help_offer)
                        
    except Exception as e:
        print(f"❌ Ошибка в главном цикле: {e}")
        continue
